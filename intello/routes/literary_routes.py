"""Literary, writing tools, reconstruction, and export routes."""
import asyncio
import json
import os
import re
import time
from typing import Optional

from fastapi import APIRouter, File, Form, Request, UploadFile
from fastapi.responses import HTMLResponse, Response

from intello import literary
from intello import gdrive
from intello import ratelimit
from intello import memory
from intello.craft import build_craft_prompt
from intello.guardrails import check_word_count
from intello import workflow as wf
from intello import writing_tools as wt
from intello import reconstruct as recon
from intello.backends import execute

router = APIRouter(tags=["literary"])


def _get_providers():
    from intello.web import _providers
    return _providers


def _get_user(request):
    from intello.web import _get_user as gu
    return gu(request) if request else "anonymous"


def _filter_providers(providers, user):
    from intello.web import filter_providers_for_user
    return filter_providers_for_user(providers, user)


# --- Literary Engine ---

@router.get("/api/literary/projects")
async def api_literary_projects():
    return literary.list_projects()


@router.post("/api/literary/projects")
async def api_literary_create_project(
    title: str = Form(...),
    genre: str = Form("fiction"),
    brief: str = Form(""),
    target_words: int = Form(0),
    style: str = Form(""),
    steps: str = Form("[]"),
):
    import uuid as _uuid
    pid = title.replace(" ", "_").lower()[:30] + f"_{int(time.time())}"
    try:
        steps_list = json.loads(steps) if steps else []
    except json.JSONDecodeError:
        steps_list = [s.strip() for s in steps.split("\n") if s.strip()]
    return literary.create_project(pid, title, genre, brief, target_words, style, steps_list)


@router.get("/api/literary/projects/{project_id}")
async def api_literary_get_project(project_id: str):
    p = literary.get_project(project_id)
    if not p:
        return {"error": "Project not found"}
    return p


@router.post("/api/literary/projects/{project_id}")
async def api_literary_update_project(
    project_id: str,
    title: Optional[str] = Form(None),
    genre: Optional[str] = Form(None),
    brief: Optional[str] = Form(None),
    target_words: Optional[int] = Form(None),
    style: Optional[str] = Form(None),
    steps: Optional[str] = Form(None),
):
    kwargs = {}
    if title is not None: kwargs["title"] = title
    if genre is not None: kwargs["genre"] = genre
    if brief is not None: kwargs["brief"] = brief
    if target_words is not None: kwargs["target_words"] = target_words
    if style is not None: kwargs["style"] = style
    if steps is not None:
        try:
            kwargs["steps"] = json.loads(steps)
        except json.JSONDecodeError:
            kwargs["steps"] = [s.strip() for s in steps.split("\n") if s.strip()]
    return literary.update_project(project_id, **kwargs)


@router.post("/api/literary/projects/{project_id}/auto-populate")
async def api_literary_auto_populate(project_id: str, doc_id: str = Form(...)):
    """Use LLMs to auto-fill project fields from the document text."""
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}

    chunks = literary.get_chunks(doc_id)
    # Send first + middle + last chunks as samples
    sample_ids = [chunks[0]["chunk_id"]] if chunks else []
    if len(chunks) > 2:
        sample_ids.append(chunks[len(chunks)//2]["chunk_id"])
    if len(chunks) > 1:
        sample_ids.append(chunks[-1]["chunk_id"])

    samples = ""
    for cid in sample_ids:
        ch = literary.get_chunk(cid)
        if ch:
            samples += ch["text"][:2000] + "\n\n"

    prompt = f"""Analyze this text and extract project metadata. Respond ONLY with valid JSON.

TEXT SAMPLES ({info['total_words']} words total):
{samples[:6000]}

Return this exact JSON structure (fill every field based on the text):
{{
  "genre": "fiction|non-fiction|screenplay|poetry|academic|technical",
  "brief": "2-3 sentence summary of what this text is about",
  "detected_style": "describe the writing style in 1-2 sentences",
  "detected_intent": "what is the author trying to achieve",
  "tone": "e.g. dark, humorous, formal, intimate, detached",
  "pov": "e.g. first person, third person limited, omniscient",
  "setting": "where and when the story takes place",
  "audience": "who is this written for",
  "themes": ["theme1", "theme2", "theme3"],
  "character_arcs": [
    {{"name": "Character Name", "arc": "brief description of their journey/role"}},
  ],
  "steps": ["major plot point or section 1", "major plot point 2", "..."],
  "target_words": estimated_final_word_count_as_integer
}}"""

    # Use a fast model for this
    for p in _providers:
        if p.available and p.provider in ("groq", "mistral"):
            result = await execute(p, prompt, max_tokens=1500,
                                   system="You are a literary analyst. Respond ONLY with valid JSON, no markdown.")
            if not result.degraded:
                try:
                    m = re.search(r'\{.*\}', result.content, re.DOTALL)
                    if m:
                        data = json.loads(m.group())
                        # Update project with extracted data
                        literary.update_project(project_id, **data)
                        return {"ok": True, "extracted": data}
                except (json.JSONDecodeError, AttributeError):
                    pass
    return {"error": "Could not auto-populate — LLM failed to return valid JSON"}


@router.get("/api/literary/workflow/{project_id}")
async def api_workflow_state(project_id: str):
    return wf.get_workflow_state(project_id)


@router.post("/api/literary/workflow/{project_id}/next")
async def api_workflow_next(
    project_id: str,
    doc_id: str = Form(""),
    mode: str = Form("horizontal"),  # horizontal or vertical
    budget_pct: int = Form(10),      # % of daily credits to spend
):
    """Execute the next logical step in the writing workflow."""
    state = wf.get_workflow_state(project_id)
    if state.get("error"):
        return state

    proj = literary.get_project(project_id)
    doc_text = literary.get_full_text(doc_id) if doc_id else ""

    # Build prompt based on mode
    if mode == "vertical":
        prompt = wf.build_vertical_prompt(proj, state, doc_text, budget_pct)
    else:
        prompt = wf.build_horizontal_prompt(proj, state, doc_text, budget_pct)

    # Inject craft techniques
    craft = build_craft_prompt(
        proj.get("genre", "fiction"),
        ["structure"] if mode == "vertical" else ["prose"],
        proj.get("style", "") or proj.get("detected_style", "")
    )
    if craft:
        prompt += "\n\n" + craft

    # Pick model based on budget — cheap for low budget, best for high
    provider = None
    if budget_pct <= 5:
        for p in _providers:
            if p.available and p.provider in ("groq", "cloudflare"):
                provider = p; break
    elif budget_pct <= 25:
        for p in _providers:
            if p.available and p.provider in ("groq", "mistral", "deepseek"):
                provider = p; break

    if not provider:
        for p in _providers:
            if p.available:
                provider = p; break

    if not provider:
        return {"error": "No providers available"}

    max_tokens = min(8192, max(1000, int(budget_pct * 80)))
    result = await execute(provider, prompt, max_tokens=max_tokens,
                           system="You are a master novelist/writer. Write with precision, depth, and craft.")

    response = {
        "state": state,
        "mode": mode,
        "budget_pct": budget_pct,
        "model": result.provider_name,
        "content": result.content if not result.degraded else f"Failed: {result.content}",
        "word_count": len(result.content.split()) if not result.degraded else 0,
        "cost": result.cost,
    }

    # If outline phase and successful, try to auto-update project
    if not result.degraded and state["phase"] == "outline":
        # Try to extract structure from the response
        try:
            fd = {"steps": [l.strip().lstrip("0123456789.-) ") for l in result.content.split("\n")
                            if l.strip() and any(l.strip().startswith(str(i)) for i in range(1, 20))]}
            if fd["steps"]:
                literary.update_project(project_id, steps=fd["steps"][:15])
        except Exception:
            pass

    # Mark step complete if we were expanding a specific step
    if not result.degraded and state["phase"] == "expand" and state["current_step_idx"] < state["steps_total"]:
        wf.mark_step_complete(project_id, state["current_step_idx"])

    # Compute next state
    response["next_state"] = wf.get_workflow_state(project_id)
    return response


@router.post("/api/literary/{doc_id}/iterate")
async def api_literary_iterate(
    doc_id: str,
    project_id: str = Form(""),
    resume: bool = Form(False),
):
    """Run iterative analysis: chunk by chunk, saving progress. Resumable."""
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}

    proj = literary.get_project(project_id) if project_id else None
    state = (proj or {}).get("iteration_state", {}) if resume else {}
    completed_chunks = state.get("completed", [])
    results_so_far = state.get("results", [])

    chunks = literary.get_chunks(doc_id)
    remaining = [c for c in chunks if c["chunk_id"] not in completed_chunks]

    if not remaining:
        return {"status": "complete", "message": "All chunks processed", "results": results_so_far}

    # Process next chunk
    chunk = remaining[0]
    chunk_data = literary.get_chunk(chunk["chunk_id"])
    if not chunk_data:
        return {"error": f"Chunk {chunk['chunk_id']} not found"}

    project_brief = literary.get_project_brief_prompt(project_id) if project_id else ""
    craft_ref = ""
    if proj:
        from intello.craft import build_craft_prompt
        craft_ref = build_craft_prompt(proj.get("genre", "fiction"), ["prose"], proj.get("style", ""))

    chunk_prompt = f"""{f"PROJECT:{chr(10)}{project_brief}{chr(10)}{chr(10)}" if project_brief else ""}{craft_ref}

Analyze this section (chunk {len(completed_chunks)+1}/{len(chunks)}):
Chapter: {chunk_data['chapter']}
Lines {chunk_data['start_line']}-{chunk_data['end_line']}

TEXT:
{chunk_data['text']}

Provide:
1. Quality assessment (1-2 sentences)
2. Specific issues found
3. Concrete edits in format: EDIT LINE X-Y: [text] — REASON: [why]
4. How this section serves the overall work

Be surgical and specific."""

    result = await execute(
        next((p for p in _providers if p.available and p.provider in ("groq", "cloudflare", "mistral")), _providers[0]),
        chunk_prompt, max_tokens=2000,
        system="You are a literary editor doing a line-by-line review. Be precise."
    )

    chunk_result = {
        "chunk_id": chunk["chunk_id"],
        "chapter": chunk["chapter"],
        "lines": f"{chunk['start_line']}-{chunk['end_line']}",
        "analysis": result.content if not result.degraded else f"Failed: {result.content}",
        "model": result.model_id,
    }

    # Parse edits
    if not result.degraded:
        for m in re.finditer(r'EDIT LINE[S]? (\d+)-(\d+):\s*(.+?)(?:\s*—\s*REASON:\s*(.+?))?(?:\n|$)', result.content):
            literary.propose_edit(doc_id, "replace", int(m.group(1)), int(m.group(2)),
                                  m.group(3).strip(), m.group(4) or "", result.model_id)

    completed_chunks.append(chunk["chunk_id"])
    results_so_far.append(chunk_result)

    # Save progress
    new_state = {"completed": completed_chunks, "results": results_so_far,
                 "total_chunks": len(chunks), "last_updated": time.time()}
    if project_id:
        literary.update_project(project_id, iteration_state=new_state)

    return {
        "status": "in_progress",
        "progress": f"{len(completed_chunks)}/{len(chunks)}",
        "chunk_result": chunk_result,
        "remaining": len(remaining) - 1,
    }


@router.post("/api/literary/ingest")
async def api_literary_ingest(
    title: str = Form(""),
    file: Optional[UploadFile] = File(None),
    text: Optional[str] = Form(None),
    gdrive_url: Optional[str] = Form(None),
    project_id: Optional[str] = Form(""),
):
    """Ingest a document for literary analysis. Supports .txt, .md, .pdf, .epub."""
    import tempfile
    fname = title
    doc_id = None

    if file and file.filename:
        fname = fname or file.filename
        ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
        doc_id = fname.replace(" ", "_").lower()[:40] + f"_{int(time.time())}"

        if ext == "pdf":
            with tempfile.NamedTemporaryFile(suffix=".pdf", delete=False) as f:
                f.write(await file.read())
                tmp = f.name
            result = literary.ingest_pdf(doc_id, tmp, fname, project_id or "")
            os.unlink(tmp)
            return result
        elif ext == "epub":
            with tempfile.NamedTemporaryFile(suffix=".epub", delete=False) as f:
                f.write(await file.read())
                tmp = f.name
            result = literary.ingest_epub(doc_id, tmp, fname, project_id or "")
            os.unlink(tmp)
            return result
        else:
            content = (await file.read()).decode("utf-8", errors="replace")
    elif text:
        content = text
        fname = fname or "pasted_text"
    elif gdrive_url and gdrive_url.strip():
        if gdrive.is_authenticated():
            content = gdrive.fetch_private(gdrive_url.strip())
        else:
            content = await gdrive.fetch_public(gdrive_url.strip())
        fname = fname or "gdrive_doc"
    else:
        return {"error": "No content provided"}

    if not content or len(content) < 50:
        return {"error": "Content too short (min 50 chars)"}

    if not doc_id:
        doc_id = fname.replace(" ", "_").lower()[:40] + f"_{int(time.time())}"
    return literary.ingest_document(doc_id, content, fname, project_id or "")


@router.get("/api/literary/documents")
async def api_literary_documents():
    with literary._db() as conn:
        rows = conn.execute("SELECT doc_id, title, total_lines, total_words, created_at FROM documents ORDER BY created_at DESC").fetchall()
    return [dict(r) for r in rows]


@router.get("/api/literary/{doc_id}")
async def api_literary_doc(doc_id: str):
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}
    return {
        "info": info,
        "structure": literary.get_structure(doc_id),
        "chunks": literary.get_chunks(doc_id),
        "pacing": literary.get_pacing_data(doc_id, window=max(5, info["total_lines"] // 20)),
        "characters": literary.get_characters(doc_id),
        "threads": literary.get_threads(doc_id),
    }


@router.get("/api/literary/{doc_id}/lines")
async def api_literary_lines(doc_id: str, start: int = 1, end: int = 50):
    return literary.get_lines(doc_id, start, end)


@router.get("/api/literary/{doc_id}/edits")
async def api_literary_edits(doc_id: str):
    return literary.get_pending_edits(doc_id)


# --- Writing Tools (AI-powered) ---

@router.post("/api/tools/transform")
async def api_tool_transform(
    text: str = Form(...),
    tool: str = Form(...),  # show_not_tell, describe, tone_shift, shrink, first_draft, brainstorm
    context: str = Form(""),
    target: str = Form(""),  # tone name, shrink format, brainstorm category
    genre: str = Form("fiction"),
    word_count: int = Form(1000),
    style: str = Form(""),
):
    """Universal writing tool endpoint."""
    prompts = {
        "show_not_tell": wt.show_not_tell(text, context),
        "describe": wt.describe_senses(text, context),
        "tone_shift": wt.tone_shift(text, target or "darker"),
        "shrink": wt.shrink_ray(text, target or "blurb"),
        "first_draft": wt.first_draft(text, style, word_count),
        "brainstorm": wt.brainstorm(text, target or "plot", genre),
    }
    prompt = prompts.get(tool)
    if not prompt:
        return {"error": f"Unknown tool: {tool}"}

    # Pick a good model
    provider = None
    for p in _providers:
        if p.available and p.provider in ("groq", "mistral", "google"):
            provider = p
            break
    if not provider:
        for p in _providers:
            if p.available:
                provider = p
                break
    if not provider:
        return {"error": "No providers available"}

    result = await execute(provider, prompt, max_tokens=4096,
                           system="You are a master fiction editor and writing coach. Be specific and craft-aware.")
    return {
        "tool": tool,
        "result": result.content if not result.degraded else f"Failed: {result.content}",
        "model": result.provider_name,
        "cost": result.cost,
        "word_count": len(result.content.split()) if not result.degraded else 0,
    }


@router.post("/api/tools/beta-read")
async def api_tool_beta_read(text: str = Form(...)):
    """Run 3 AI beta readers in parallel with different perspectives."""
    import asyncio
    reader_types = ["casual", "craft", "market"]

    # Pick 3 different models for diversity
    models = []
    seen = set()
    for p in _providers:
        if p.available and p.provider not in seen and ratelimit.remaining(p.model_id, p.daily_limit) != 0:
            models.append(p)
            seen.add(p.provider)
            if len(models) >= 3:
                break

    if not models:
        return {"error": "Not enough providers available"}

    tasks = []
    for i, rtype in enumerate(reader_types):
        model = models[i % len(models)]
        prompt = wt.beta_reader_prompt(text, rtype)
        tasks.append(execute(model, prompt, max_tokens=2000,
                             system=f"You are a {rtype} reader. Give honest, detailed feedback."))

    results = await asyncio.gather(*tasks)
    readers = []
    for rtype, result in zip(reader_types, results):
        readers.append({
            "type": rtype,
            "model": result.provider_name,
            "feedback": result.content if not result.degraded else "Failed",
            "cost": result.cost,
        })

    return {"readers": readers, "total_cost": sum(r["cost"] for r in readers)}


@router.post("/api/literary/{doc_id}/analyze")
async def api_literary_analyze(doc_id: str, focus: str = Form("full"),
                                async_mode: bool = Form(False)):
    """Run multi-model literary analysis. Add async_mode=true for background processing."""
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}

    structure = literary.get_structure(doc_id)
    pacing = literary.get_pacing_data(doc_id, window=30)
    chunks = literary.get_chunks(doc_id)

    # Build analysis prompt with structure + pacing summary
    struct_summary = "\n".join(
        f"  {json.loads(s['metadata']).get('title','?')} (lines {s['start_line']}-{s['end_line']})"
        for s in structure
    )
    pacing_summary = ""
    for p in pacing:
        tension_label = "🔴 HIGH" if p["tension"] > 1.5 else "🟡 MED" if p["tension"] > 0.5 else "🟢 LOW"
        pacing_summary += f"  Lines {p['start_line']}-{p['end_line']}: tension={tension_label}, dialogue={p['dialogue_ratio']:.0%}, avg_sent={p['avg_sentence_len']:.1f} words\n"

    # Pick 2-3 chunks to send as samples (beginning, middle, end)
    sample_ids = []
    if chunks:
        sample_ids = [chunks[0]["chunk_id"]]
        if len(chunks) > 2:
            sample_ids.append(chunks[len(chunks) // 2]["chunk_id"])
        if len(chunks) > 1:
            sample_ids.append(chunks[-1]["chunk_id"])

    samples = ""
    for cid in sample_ids:
        ch = literary.get_chunk(cid)
        if ch:
            samples += f"\n--- {ch['chapter']} (lines {ch['start_line']}-{ch['end_line']}) ---\n{ch['text'][:3000]}\n"

    # Get project brief if linked
    project_brief = ""
    project_genre = "fiction"
    project_style = ""
    project_target = 0
    doc_project_id = info.get("project_id") or ""
    if isinstance(doc_project_id, str) and doc_project_id:
        project_brief = literary.get_project_brief_prompt(doc_project_id)
        proj = literary.get_project(doc_project_id)
        if proj:
            project_genre = proj.get("genre", "fiction")
            project_style = proj.get("style", "")
            project_target = proj.get("target_words", 0)

    # Detect issues from pacing data for craft reference
    pacing_issues = []
    for p in pacing:
        if p["tension"] < 0.3:
            pacing_issues.append("slow pacing")
        if p["tension"] > 2:
            pacing_issues.append("fast pacing")
        if p["dialogue_ratio"] < 0.05:
            pacing_issues.append("flat — no dialogue")
    pacing_issues = list(set(pacing_issues))[:5]

    craft_ref = build_craft_prompt(project_genre, pacing_issues, project_style)

    # Word count check
    wc_note = ""
    if project_target:
        wc = check_word_count(literary.get_full_text(doc_id), project_target)
        wc_note = f"\nWORD COUNT: {wc['actual']} / {wc['target']} target ({wc['verdict']})"

    analysis_prompt = f"""Analyze this document as a literary expert.

{f"PROJECT CONTEXT:{chr(10)}{project_brief}{chr(10)}" if project_brief else ""}{wc_note}

DOCUMENT: "{info['title']}" — {info['total_lines']} lines, {info['total_words']} words

STRUCTURE:
{struct_summary}

PACING ANALYSIS:
{pacing_summary}

{craft_ref}

SAMPLE EXCERPTS:
{samples}

Provide a detailed literary analysis covering:
1. STRUCTURE: Is the chapter/scene organization effective? Where should breaks be added/removed?
2. PACING: Where is the narrative too slow or too fast? Which sections need tightening or expansion?
3. PROSE QUALITY: Comment on sentence variety, word choice, rhythm. Quote specific lines.
4. WORD COUNT: Is the current length appropriate for the target? Which sections are bloated or too thin?
5. SUGGESTIONS: For each issue, provide SPECIFIC edits in this format:
   - EDIT LINE X-Y: [replacement text] — REASON: [why]
   - INSERT AFTER LINE X: [new text] — REASON: [why]

IMPORTANT: When you suggest new or replacement text, it MUST hit the word count implied by the edit.
Do NOT write "a 500-word paragraph about X" — actually WRITE the 500 words.

Be brutally honest. This writer wants to produce super literature, not hear compliments."""

    # Async mode — return job_id
    if async_mode:
        job_id = jobsys.create_job("literary_analyze", f"Analyze {info['title']}")

        async def _run_analysis():
            from intello.pipeline import run_deep as _rd
            pipe = await _rd(analysis_prompt, _get_providers())
            edits = 0
            if pipe.final and not pipe.final.degraded:
                for m in re.finditer(r'EDIT LINE[S]? (\d+)-(\d+):\s*(.+?)(?:\s*—\s*REASON:\s*(.+?))?(?:\n|$)', pipe.final.content):
                    literary.propose_edit(doc_id, "replace", int(m.group(1)), int(m.group(2)),
                                          m.group(3).strip(), m.group(4) or "", pipe.final.model_id)
                    edits += 1
            return {"analysis": pipe.final.content if pipe.final else "Failed",
                    "edits_proposed": edits, "cost": pipe.total_cost}

        asyncio.create_task(jobsys.run_async(job_id, _run_analysis()))
        return {"job_id": job_id, "status": "queued",
                "poll": f"/api/jobs/{job_id}", "result": f"/api/jobs/{job_id}/result"}

    # Sync mode
    from intello.pipeline import run_deep
    pipe = await run_deep(analysis_prompt, _get_providers())

    # Word count verification of the analysis output
    wc_result = None
    if project_target and pipe.final and not pipe.final.degraded:
        wc_result = check_word_count(literary.get_full_text(doc_id), project_target)

    # Parse any edit suggestions from the response
    edits_proposed = 0
    if pipe.final and not pipe.final.degraded:
        # Try to extract EDIT LINE patterns
        for m in re.finditer(r'EDIT LINE[S]? (\d+)-(\d+):\s*(.+?)(?:\s*—\s*REASON:\s*(.+?))?(?:\n|$)',
                             pipe.final.content):
            literary.propose_edit(doc_id, "replace", int(m.group(1)), int(m.group(2)),
                                  m.group(3).strip(), m.group(4) or "", pipe.final.model_id)
            edits_proposed += 1
        for m in re.finditer(r'INSERT AFTER LINE (\d+):\s*(.+?)(?:\s*—\s*REASON:\s*(.+?))?(?:\n|$)',
                             pipe.final.content):
            literary.propose_edit(doc_id, "insert", int(m.group(1)), int(m.group(1)),
                                  m.group(2).strip(), m.group(3) or "", pipe.final.model_id)
            edits_proposed += 1

    return {
        "analysis": pipe.final.content if pipe.final else "Analysis failed",
        "edits_proposed": edits_proposed,
        "pipeline_steps": pipe.steps_log,
        "cost": pipe.total_cost,
        "word_count": wc_result,
        "craft_techniques_used": len(pacing_issues),
    }


@router.post("/api/literary/{doc_id}/edit/{edit_id}/apply")
async def api_literary_apply_edit(doc_id: str, edit_id: int):
    ok = literary.apply_edit(edit_id)
    return {"ok": ok}


@router.post("/api/literary/{doc_id}/edit/{edit_id}/reject")
async def api_literary_reject_edit(doc_id: str, edit_id: int):
    literary.reject_edit(edit_id)
    return {"ok": True}


@router.post("/api/literary/{doc_id}/append")
async def api_literary_append(doc_id: str, text: str = Form(...)):
    """Append text to a document (from workflow output, etc.)."""
    info = literary.get_document_info(doc_id)
    if not info:
        return {"error": "Document not found"}
    current = literary.get_full_text(doc_id)
    new_text = current + "\n\n" + text
    # Re-ingest with appended content
    result = literary.ingest_document(doc_id, new_text, info["title"],
                                      info.get("project_id", ""))
    return result


@router.get("/api/literary/{doc_id}/export/docx")
async def api_literary_export_docx(doc_id: str):
    """Export document as DOCX."""
    info = literary.get_document_info(doc_id)
    if not info:
        return Response("Not found", status_code=404)

    from docx import Document as DocxDocument
    from docx.shared import Pt
    import tempfile

    doc = DocxDocument()
    doc.add_heading(info["title"], 0)

    with literary._db() as conn:
        lines = conn.execute("SELECT line_num, text, chapter FROM lines WHERE doc_id=? ORDER BY line_num",
                             (doc_id,)).fetchall()

    current_chapter = ""
    for line in lines:
        if line["chapter"] != current_chapter:
            current_chapter = line["chapter"]
            doc.add_heading(current_chapter, level=1)
        elif line["text"].strip():
            doc.add_paragraph(line["text"])

    with tempfile.NamedTemporaryFile(suffix=".docx", delete=False) as f:
        doc.save(f.name)
        tmp = f.name

    with open(tmp, "rb") as f:
        content = f.read()
    os.unlink(tmp)

    return Response(content, media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    headers={"Content-Disposition": f'attachment; filename="{info["title"]}.docx"'})


@router.get("/api/literary/{doc_id}/export", response_class=HTMLResponse)
async def api_literary_export(doc_id: str):
    """Generate a rich, editable HTML report of the full literary analysis."""
    info = literary.get_document_info(doc_id)
    if not info:
        return HTMLResponse("<h1>Document not found</h1>", status_code=404)

    structure = literary.get_structure(doc_id)
    characters = literary.get_characters(doc_id)
    pacing = literary.get_pacing_data(doc_id, window=max(5, info["total_lines"] // 25))
    threads = literary.get_threads(doc_id)
    edits = literary.get_pending_edits(doc_id)

    with literary._db() as conn:
        lines = conn.execute("SELECT line_num, text, chapter FROM lines WHERE doc_id=? ORDER BY line_num",
                             (doc_id,)).fetchall()
        pacing_raw = conn.execute("SELECT line_num, tension, dialogue FROM pacing WHERE doc_id=? ORDER BY line_num",
                                  (doc_id,)).fetchall()

    # Build edit lookup by line
    edit_map = {}
    for e in edits:
        for ln in range(e["start_line"], e["end_line"] + 1):
            edit_map[ln] = e

    # Pacing lookup
    pacing_map = {r["line_num"]: dict(r) for r in pacing_raw}

    # Thread color map
    thread_colors = {
        "identity": "#ef4444", "motive": "#f97316", "event": "#eab308", "method": "#22c55e",
        "location": "#06b6d4", "timing": "#3b82f6", "suspense": "#8b5cf6", "secret": "#ec4899",
        "mystery": "#a855f7", "concealment": "#6366f1", "deception": "#f43f5e", "curiosity": "#14b8a6",
        "anomaly": "#f59e0b", "decision": "#10b981", "unknown": "#6b7280", "uncertainty": "#9ca3af",
        "threat": "#dc2626", "promise": "#2563eb", "investigation": "#7c3aed",
    }

    # --- Build HTML ---
    import html as html_mod

    def esc(s):
        return html_mod.escape(str(s))

    # Pacing SVG
    max_t = max((p["tension"] for p in pacing), default=1) or 1
    svg_w, svg_h = 700, 80
    points = []
    for i, p in enumerate(pacing):
        x = (i / max(len(pacing) - 1, 1)) * svg_w
        y = svg_h - (p["tension"] / max_t) * (svg_h - 10)
        points.append(f"{x:.0f},{y:.0f}")
    pacing_svg = (
        f'<svg width="{svg_w}" height="{svg_h}" style="width:100%;height:{svg_h}px">'
        f'<polyline points="{" ".join(points)}" fill="none" stroke="#6366f1" stroke-width="2"/>'
    )
    # Add colored dots
    for i, p in enumerate(pacing):
        x = (i / max(len(pacing) - 1, 1)) * svg_w
        y = svg_h - (p["tension"] / max_t) * (svg_h - 10)
        color = "#ef4444" if p["tension"] > max_t * 0.7 else "#eab308" if p["tension"] > max_t * 0.3 else "#22c55e"
        pacing_svg += f'<circle cx="{x:.0f}" cy="{y:.0f}" r="4" fill="{color}"><title>L{p["start_line"]}-{p["end_line"]} tension={p["tension"]:.1f}</title></circle>'
    pacing_svg += '</svg>'

    # Thread bars SVG
    total = info["total_lines"]
    thread_svg = f'<svg width="100%" height="{len(threads) * 14 + 4}" style="width:100%">'
    for i, t in enumerate(threads):
        x1 = (t["start_line"] / total) * 100
        w = max(1, ((t["end_line"] - t["start_line"]) / total) * 100)
        color = thread_colors.get(t["category"], "#6b7280")
        opacity = "0.5" if t["resolved"] else "0.9"
        status = "✅" if t["resolved"] else "❓"
        thread_svg += (f'<rect x="{x1:.1f}%" y="{i * 14}" width="{w:.1f}%" height="10" rx="3" '
                       f'fill="{color}" opacity="{opacity}">'
                       f'<title>{status} {esc(t["category"])}: {esc(t["description"][:80])}</title></rect>')
    thread_svg += '</svg>'

    # Annotated text
    text_html = ""
    current_chapter = ""
    for row in lines:
        ln = row["line_num"]
        txt = esc(row["text"])
        p = pacing_map.get(ln, {})
        tension = p.get("tension", 0) if p else 0
        is_dialogue = p.get("dialogue", 0) if p else 0

        # Chapter header
        if row["chapter"] != current_chapter:
            current_chapter = row["chapter"]
            text_html += f'<h3 style="color:#6366f1;margin:24px 0 8px;page-break-before:auto" id="line-{ln}">{txt}</h3>\n'
            continue

        # Line styling
        style = ""
        cls = ""
        if tension > 1.5:
            style = "border-left:3px solid #ef4444;padding-left:8px;"
            cls = "high-tension"
        elif is_dialogue:
            style = "color:#3b82f6;"

        # Edit annotation
        edit_note = ""
        if ln in edit_map:
            e = edit_map[ln]
            edit_note = (f'<span style="background:#22c55e22;border:1px solid #22c55e;border-radius:4px;'
                         f'padding:2px 6px;font-size:.8rem;margin-left:8px" contenteditable="false">'
                         f'✏️ {esc(e["reason"][:60])}</span>')

        text_html += (f'<div style="display:flex;gap:12px;{style}" id="line-{ln}">'
                      f'<span style="color:#71717a;font-size:.75rem;min-width:35px;text-align:right;'
                      f'font-family:monospace;user-select:none">{ln}</span>'
                      f'<span contenteditable="true" style="flex:1">{txt}</span>'
                      f'{edit_note}</div>\n')

    # Character summary
    char_html = "".join(
        f'<span style="display:inline-block;background:#1a1d27;border:1px solid #2a2d3a;'
        f'padding:3px 10px;border-radius:12px;margin:3px;font-size:.85rem">'
        f'{esc(c["name"])} <span style="color:#71717a;font-size:.75rem">{c["mentions"]}×</span></span>'
        for c in characters
    )

    # Thread descriptions
    thread_desc_html = ""
    for t in threads:
        color = thread_colors.get(t["category"], "#6b7280")
        status = "✅ Resolved" if t["resolved"] else "❓ Open"
        thread_desc_html += (
            f'<div style="border-left:4px solid {color};padding:6px 12px;margin:4px 0;'
            f'background:#0f1117;border-radius:0 6px 6px 0;font-size:.85rem">'
            f'<span style="color:{color};font-size:.7rem;text-transform:uppercase;font-weight:600">{esc(t["category"])}</span> '
            f'<span style="color:#71717a;font-size:.75rem">L{t["start_line"]}–{t["end_line"]}</span> '
            f'<span style="font-size:.75rem">{status}</span><br>'
            f'{esc(t["description"])}</div>'
        )

    # Assemble
    report = f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8">
<title>{esc(info['title'])} — Literary Analysis</title>
<style>
  @media print {{ @page {{ margin: 2cm; }} .no-print {{ display: none; }} }}
  body {{ font-family: Georgia, serif; max-width: 900px; margin: 0 auto; padding: 24px;
          background: #fff; color: #1a1a1a; line-height: 1.8; }}
  h1 {{ font-size: 1.8rem; border-bottom: 2px solid #6366f1; padding-bottom: 8px; }}
  h2 {{ font-size: 1.2rem; color: #6366f1; margin: 32px 0 12px; }}
  .stats {{ color: #666; font-size: .9rem; margin-bottom: 24px; }}
  .section {{ margin: 24px 0; padding: 16px; background: #f8f9fa; border-radius: 8px; }}
  [contenteditable=true]:focus {{ outline: 2px solid #6366f1; border-radius: 4px; }}
  [contenteditable=true]:hover {{ background: #f0f0ff; }}
  .toolbar {{ position: sticky; top: 0; background: #fff; padding: 8px 0; border-bottom: 1px solid #ddd;
              z-index: 10; display: flex; gap: 8px; }}
  .toolbar button {{ padding: 6px 14px; border: 1px solid #ddd; border-radius: 6px; cursor: pointer;
                     background: #fff; font-size: .85rem; }}
  .toolbar button:hover {{ background: #f0f0ff; }}
</style>
</head>
<body>

<div class="toolbar no-print">
  <button onclick="window.print()">🖨️ Print / PDF</button>
  <button onclick="downloadHTML()">💾 Save HTML</button>
  <span style="color:#999;font-size:.85rem;margin-left:auto">Click any text to edit directly</span>
</div>

<h1>{esc(info['title'])}</h1>
<div class="stats">{info['total_lines']} lines · {info['total_words']} words · {len(structure)} chapters · {len(characters)} characters · {len(threads)} narrative threads</div>

<h2>📖 Structure</h2>
<div class="section">
{"".join(f'<div><a href="#line-{s["start_line"]}" style="color:#6366f1">{esc(json.loads(s["metadata"])["title"])}</a> <span style="color:#999;font-size:.85rem">lines {s["start_line"]}–{s["end_line"]}</span></div>' for s in structure)}
</div>

<h2>👤 Characters</h2>
<div class="section">{char_html or '<span style="color:#999">No characters detected</span>'}</div>

<h2>📊 Pacing Curve</h2>
<div class="section">
  <div style="display:flex;justify-content:space-between;font-size:.75rem;color:#999"><span>Start</span><span>🟢 Low — 🟡 Medium — 🔴 High tension</span><span>End</span></div>
  {pacing_svg}
</div>

<h2>🧵 Narrative Threads</h2>
<div class="section">
  {thread_svg}
  <div style="margin-top:12px">{thread_desc_html}</div>
</div>

<h2>📝 Annotated Text</h2>
<div style="font-size:.95rem;line-height:1.9">
{text_html}
</div>

<script>
function downloadHTML() {{
  const html = document.documentElement.outerHTML;
  const blob = new Blob([html], {{type: 'text/html'}});
  const a = document.createElement('a');
  a.href = URL.createObjectURL(blob);
  a.download = '{esc(info["title"]).replace("'", "")}_analysis.html';
  a.click();
}}
</script>
</body>
</html>"""

    return HTMLResponse(report)


